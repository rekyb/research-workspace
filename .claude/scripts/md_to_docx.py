#!/usr/bin/env python3
"""Convert a research Markdown file to .docx using python-docx.

Pragmatic converter tuned for SYNTHESIS.md / README.md structure. Handles:
  - # / ## / ### / #### headings
  - unordered (- , * ) and ordered (1. ) list items, with **bold** lead-ins
  - **bold** and *italic* inline spans, plus `code` spans
  - horizontal rules (---) and blank-line paragraph breaks
  - simple pipe tables (| a | b |)
  - embedded images (![alt](path)) resolved relative to input markdown directory
  - GitHub-style alerts and blockquotes (> [!NOTE], > [Principal Researcher]...)
  - fenced code blocks (``` ... ```) rendered cleanly without list collision

Usage: python3 md_to_docx.py INPUT.md [OUTPUT.docx]
       (defaults OUTPUT to INPUT with .docx extension)

Styling: clean white document — Calibri 12pt body, grayscale only (no accent
colours). Headings are dark grey, tables use plain black-bordered grid with a bold
header row, blockquotes are a subtle grey italic indent, code is Consolas. Colour is
avoided unless the source markdown explicitly asks for it.
"""
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.exit("python-docx is not installed. Run: pip install python-docx")

INLINE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`|!\[.+?\]\(.+?\))")
IMG_PATTERN = re.compile(r"!\[(.*?)\]\((.*?)\)")

# --- Clean, grayscale styling knobs -----------------------------------------
BODY_FONT = "Calibri"
BODY_SIZE = 12          # pt
MONO_FONT = "Consolas"
CODE_SIZE = 10.5        # pt (slightly smaller than body)
CAPTION_SIZE = 9.5      # pt
LINE_SPACING = 1.5      # body line height
INK = RGBColor(0x00, 0x00, 0x00)        # body text: black
HEADING_INK = RGBColor(0x26, 0x26, 0x26)  # headings: near-black grey (not a colour)
MUTED_INK = RGBColor(0x59, 0x59, 0x59)    # captions / quotes: medium grey
TABLE_STYLE = "Table Grid"              # plain black borders, white cells — no accent
# Heading sizes (pt) by level 1..4
HEADING_SIZES = {1: 18, 2: 15, 3: 13, 4: 12}


def apply_clean_styles(doc):
    """Force a clean, white, grayscale look: Calibri 12pt body, dark-grey headings,
    no theme accent colours."""
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_SIZE)
    normal.font.color.rgb = INK
    normal.paragraph_format.line_spacing = LINE_SPACING

    for level, size in HEADING_SIZES.items():
        try:
            st = doc.styles[f"Heading {level}"]
        except KeyError:
            continue
        st.font.name = BODY_FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.italic = False
        st.font.color.rgb = HEADING_INK  # override the default blue accent


def add_inline(paragraph, text, md_parent_dir=None, doc=None):
    """Add text to a paragraph, rendering **bold**, *italic*, `code` spans, and inline images."""
    for part in INLINE.split(text):
        if not part:
            continue
        
        img_match = IMG_PATTERN.fullmatch(part)
        if img_match:
            alt_text = img_match.group(1).strip()
            img_rel_path = img_match.group(2).strip()
            # Try to resolve and embed image if doc and md_parent_dir provided
            if doc and md_parent_dir:
                img_path = (md_parent_dir / img_rel_path).resolve()
                if img_path.is_file() and img_path.suffix.lower() in [".png", ".jpg", ".jpeg"]:
                    try:
                        p_img = doc.add_paragraph()
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_img.add_run().add_picture(str(img_path), width=Inches(5.5))
                        if alt_text:
                            p_cap = doc.add_paragraph()
                            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            r_cap = p_cap.add_run(f"Figure: {alt_text}")
                            r_cap.italic = True
                            r_cap.font.size = Pt(CAPTION_SIZE)
                            r_cap.font.color.rgb = MUTED_INK
                        continue
                    except Exception:
                        pass
            # Fallback if image cannot be embedded or is GIF/external
            run = paragraph.add_run(f"[Evidence Capture: {alt_text or img_rel_path} ({img_rel_path})]")
            run.italic = True
            continue

        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = MONO_FONT
        else:
            paragraph.add_run(part)


def flush_table(doc, rows, md_parent_dir=None):
    """Render collected pipe-table rows (list of cell-lists)."""
    rows = [r for r in rows if not re.match(r"^[\s|:-]+$", " ".join(r))]
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=cols)
    try:
        table.style = TABLE_STYLE  # plain black grid, white cells — no accent colour
    except KeyError:
        pass
    for row_idx, r in enumerate(rows):
        cells = table.add_row().cells
        for i in range(cols):
            text = r[i] if i < len(r) else ""
            cells[i].text = ""
            add_inline(cells[i].paragraphs[0], text, md_parent_dir, doc)
            if row_idx == 0:  # header row: bold, no fill (clean, colourless emphasis)
                for run in cells[i].paragraphs[0].runs:
                    run.bold = True


def convert(md_path, docx_path):
    md_file = Path(md_path)
    md_parent_dir = md_file.parent
    lines = md_file.read_text(encoding="utf-8").splitlines()
    doc = Document()
    apply_clean_styles(doc)

    table_buf = []
    in_code_block = False
    code_buf = []

    for raw in lines:
        line = raw.rstrip()

        # Check for fenced code block toggle
        if re.match(r"^\s*```", line):
            if in_code_block:
                in_code_block = False
                if code_buf:
                    p_code = doc.add_paragraph()
                    r_code = p_code.add_run("\n".join(code_buf))
                    r_code.font.name = MONO_FONT
                    r_code.font.size = Pt(CODE_SIZE)
                    code_buf = []
            else:
                in_code_block = True
                if table_buf:
                    flush_table(doc, table_buf, md_parent_dir)
                    table_buf = []
            continue

        if in_code_block:
            code_buf.append(raw)
            continue

        if line.startswith("|") and line.count("|") >= 2:
            table_buf.append([c.strip() for c in line.strip("|").split("|")])
            continue
        elif table_buf:
            flush_table(doc, table_buf, md_parent_dir)
            table_buf = []

        if not line.strip():
            continue
        if re.match(r"^-{3,}$", line.strip()):
            continue

        # Check for standalone image line
        m_img = IMG_PATTERN.fullmatch(line.strip())
        if m_img:
            p_dummy = doc.add_paragraph()
            add_inline(p_dummy, line.strip(), md_parent_dir, doc)
            if not p_dummy.text.strip():
                # If picture added cleanly, remove empty paragraph
                p_element = p_dummy._element
                p_element.getparent().remove(p_element)
            continue

        # Check for blockquotes / alerts
        m_quote = re.match(r"^\s*>\s*(.*)", line)
        if m_quote:
            quote_text = m_quote.group(1).strip()
            # Clean, colourless blockquote: indented grey italic (no accent box).
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_inline(p, quote_text, md_parent_dir, doc)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = MUTED_INK
            continue

        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            doc.add_heading(m.group(2).strip(), level=len(m.group(1)))
            continue

        m = re.match(r"^\s*[-*]\s+(.*)", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, m.group(1), md_parent_dir, doc)
            continue

        m = re.match(r"^\s*\d+\.\s+(.*)", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, m.group(1), md_parent_dir, doc)
            continue

        p = doc.add_paragraph()
        add_inline(p, line, md_parent_dir, doc)

    if table_buf:
        flush_table(doc, table_buf, md_parent_dir)
    if code_buf:
        p_code = doc.add_paragraph()
        r_code = p_code.add_run("\n".join(code_buf))
        r_code.font.name = "Consolas"
        r_code.font.size = Pt(9.5)

    doc.save(docx_path)
    return docx_path


if __name__ == "__main__":
    if len(sys.argv) < 2:
        sys.exit("Usage: python3 md_to_docx.py INPUT.md [OUTPUT.docx]")
    src = sys.argv[1]
    dst = sys.argv[2] if len(sys.argv) > 2 else str(Path(src).with_suffix(".docx"))
    print(convert(src, dst))
